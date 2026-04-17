#!/usr/bin/env python3
"""
tracked_change_generator.py
===========================
Generates manuscript with tracked (marked) changes for peer review revision.

Creates:
  1. Tracked changes version: shows deletions (strikethrough) + insertions (underline)
  2. Clean revised version: the revised manuscript without markup
  3. Change manifest: list of all changes with locations and rationales

Supports two output modes:
  - markdown: Uses ~~deleted~~ and __inserted__ markup (human-readable)
  - docx:    Uses python-docx revision marks (submission-ready)

Usage (import as module):
    from tracked_change_generator import TrackedChangeGenerator
    gen = TrackedChangeGenerator()
    changes = [
        {"location": "Page 3, Para 1", "old": "...", "new": "..."},
        {"location": "Page 5, Para 2", "old": "...", "new": "..."},
    ]
    tracked = gen.generate_tracked("manuscript_original.md", changes)
    clean = gen.generate_clean("manuscript_original.md", changes)

CLI usage:
    python tracked_change_generator.py manuscript.md --changes changes.json -o revised/
"""

from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class Change:
    """
    Represents a single text change.

    Fields:
        location:  Manuscript location (e.g. "Page 3, Para 2")
        old:       Original text (exact)
        new:       Revised text
        comment_id: Link to reviewer comment ID (e.g. "R1-C1")
        rationale:  Why this change was made
        type:      text_revision | addition | deletion | substitution
        section:   Which manuscript section (Introduction/Methods/Results/Discussion)
        status:    pending | applied
    """

    location: str
    old: str
    new: str
    comment_id: str = ""
    rationale: str = ""
    type: str = "substitution"
    section: str = ""
    status: str = "pending"


# ─── Core Change Engine ────────────────────────────────────────────────────────


class TrackedChangeGenerator:
    """
    Generates manuscripts with tracked changes.

    Strategy: Find exact text matches in the manuscript, replace with tracked markup.
    Falls back to fuzzy matching when exact match fails.
    """

    def __init__(self, tolerance: float = 0.8):
        """
        Args:
            tolerance: Similarity threshold for fuzzy matching (0.0-1.0).
                      1.0 = exact match only. Lower = more aggressive matching.
        """
        self.tolerance = tolerance
        self._changes_applied: list[Change] = []

    def _find_match(
        self, text: str, old: str
    ) -> tuple[int, int] | None:
        """
        Find the position of 'old' text within 'text'.

        Returns (start, end) indices or None if not found.
        Uses exact match first, then fuzzy fallback.
        """
        # Exact match
        idx = text.find(old)
        if idx >= 0:
            return idx, idx + len(old)

        # Fuzzy match
        old_words = old.split()
        if len(old_words) < 3:
            return None  # Too short for fuzzy

        # Try to find the old text by searching for key words
        key_words = [w for w in old_words if len(w) > 4]
        if not key_words:
            return None

        # Find a sentence containing the most key words
        best_match = None
        best_score = 0

        # Split into sentences
        sentences = re.split(r"(?<=[.!?])\s+", text)
        pos = 0
        for sentence in sentences:
            score = sum(1 for kw in key_words if kw.lower() in sentence.lower())
            if score > best_score:
                best_score = score
                best_match = (pos, pos + len(sentence))
            pos += len(sentence) + 1

        if best_match and best_score >= len(key_words) * self.tolerance:
            return best_match

        return None

    def _apply_tracked(
        self, text: str, change: Change
    ) -> tuple[str, bool]:
        """
        Apply a single tracked change to text.

        Returns (new_text, success).
        """
        match = self._find_match(text, change.old)
        if match is None:
            return text, False

        start, end = match
        tracked = f"~~{change.old}~~**→**{change.new}"

        return text[:start] + tracked + text[end:], True

    def _apply_clean(
        self, text: str, change: Change
    ) -> tuple[str, bool]:
        """Apply a clean change (no markup) to text."""
        match = self._find_match(text, change.old)
        if match is None:
            return text, False

        start, end = match
        return text[:start] + change.new + text[end:], True

    # ─── Main generation methods ───────────────────────────────────────────

    def generate_tracked(
        self,
        manuscript_path: str | Path,
        changes: list[dict[str, str] | Change],
        output_path: str | Path | None = None,
    ) -> str:
        """
        Generate the tracked-changes version of a manuscript.

        Args:
            manuscript_path: Path to the original manuscript markdown
            changes:         List of Change dicts or Change objects
            output_path:      Optional output file path

        Returns:
            The tracked manuscript as a string.
        """
        text = Path(manuscript_path).read_text(encoding="utf-8")
        self._changes_applied = []

        for c_data in changes:
            if isinstance(c_data, dict):
                c = Change(**c_data)
            else:
                c = c_data

            new_text, success = self._apply_tracked(text, c)
            if success:
                c.status = "applied"
                self._changes_applied.append(c)
            else:
                c.status = "failed"
            text = new_text

        if output_path:
            Path(output_path).write_text(text, encoding="utf-8")

        return text

    def generate_clean(
        self,
        manuscript_path: str | Path,
        changes: list[dict[str, str] | Change],
        output_path: str | Path | None = None,
    ) -> str:
        """
        Generate the clean revised manuscript (no markup).

        Args:
            manuscript_path: Path to the original manuscript markdown
            changes:         List of Change dicts or Change objects
            output_path:      Optional output file path

        Returns:
            The clean revised manuscript as a string.
        """
        text = Path(manuscript_path).read_text(encoding="utf-8")

        for c_data in changes:
            if isinstance(c_data, dict):
                c = Change(**c_data)
            else:
                c = c_data

            text, success = self._apply_clean(text, c)
            if success:
                c.status = "applied"
            else:
                c.status = "failed"

        if output_path:
            Path(output_path).write_text(text, encoding="utf-8")

        return text

    def generate_manifest(
        self,
        changes: list[dict[str, str] | Change],
        output_path: str | Path | None = None,
    ) -> str:
        """
        Generate a change manifest listing all changes with locations and rationales.

        Returns:
            Markdown table of all changes.
        """
        lines = [
            "# Revision Change Manifest\n",
            "| # | Location | Section | Type | Comment | Rationale | Status |",
            "|--:|----------|---------|------|---------|-----------|--------|",
        ]

        for i, c_data in enumerate(changes, 1):
            if isinstance(c_data, dict):
                c = Change(**c_data)
            else:
                c = c_data

            # Truncate old/new for display
            old_short = c.old[:40] + "..." if len(c.old) > 40 else c.old
            new_short = c.new[:40] + "..." if len(c.new) > 40 else c.new

            lines.append(
                f"| {i} | {c.location} | {c.section or '—'} | "
                f"{c.type} | {c.comment_id or '—'} | "
                f"{c.rationale or '—'} | {c.status} |"
            )
            lines.append(
                f"| | ~~{old_short}~~ | | | | | |"
            )
            lines.append(
                f"| | __{new_short}__ | | | | | |"
            )

        # Summary
        applied = sum(1 for c_data in changes
                      if (c_data.get("status") if isinstance(c_data, dict) else c_data.status) == "applied")
        failed = sum(1 for c_data in changes
                     if (c_data.get("status") if isinstance(c_data, dict) else c_data.status) == "failed")
        total = len(changes)

        lines.extend([
            "",
            "## Summary\n",
            f"- **Total changes:** {total}",
            f"- **Applied:** {applied} ({applied/total*100:.0f}%)",
            f"- **Failed:** {failed} ({failed/total*100:.0f}%)",
            "",
            "## Failed Changes\n",
        ])

        for i, c_data in enumerate(changes, 1):
            c = Change(**c_data) if isinstance(c_data, dict) else c_data
            if c.status == "failed":
                lines.append(
                    f"- ⚠️ **{i}. {c.comment_id or 'Unknown'}:** "
                    f"{c.location} — could not locate: \"{c.old[:60]}\"\n"
                )

        manifest = "\n".join(lines)
        if output_path:
            Path(output_path).write_text(manifest, encoding="utf-8")
        return manifest

    def export_docx(
        self,
        manuscript_path: str | Path,
        changes: list[dict[str, str] | Change],
        output_path: str | Path,
        mode: str = "clean",
    ) -> None:
        """
        Export tracked changes to .docx.

        Args:
            manuscript_path: Path to the original manuscript
            changes:          List of changes
            output_path:      Output .docx path
            mode:             "clean" (no markup) or "tracked" (with revision marks)

        Note:
            Full revision mark support in python-docx is limited.
            This generates a clean .docx with annotations in comments.
            For proper tracked changes, use Word's Compare Documents feature.
        """
        try:
            from docx import Document
            from docx.shared import RGBColor
            from docx.enum.text import WD_COLOR_INDEX
        except ImportError:
            sys.stderr.write(
                "python-docx not available. Install with: pip install python-docx\n"
                "Falling back to markdown output.\n"
            )
            # Fallback: output clean markdown
            text = self.generate_clean(manuscript_path, changes)
            fallback = str(output_path).replace(".docx", ".md")
            Path(fallback).write_text(text, encoding="utf-8")
            return

        text = Path(manuscript_path).read_text(encoding="utf-8")

        # For tracked mode: insert change markers as comments or special formatting
        # Since python-docx doesn't support true tracked changes natively,
        # we add inline markers and use a comment annotation
        if mode == "tracked":
            # Apply tracked markup
            for c_data in changes:
                c = Change(**c_data) if isinstance(c_data, dict) else c_data
                new_text, success = self._apply_tracked(text, c)
                if success:
                    c.status = "applied"
                else:
                    c.status = "failed"
                text = new_text

        doc = Document()
        doc.add_heading("Revised Manuscript", 0)

        # Add change summary at top
        applied = sum(
            1 for c_data in changes
            if (c_data.get("status") if isinstance(c_data, dict) else c_data.status) == "applied"
        )
        p = doc.add_paragraph()
        p.add_run(f"Total changes applied: {applied}/{len(changes)}\n")
        p.add_run("For full tracked changes, use Word's Compare Documents feature on the original and revised files.\n")

        # Parse markdown and write to docx
        lines = text.splitlines()
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                doc.add_paragraph()
                continue

            # Skip manifest markers
            if "~~" in line or "**→**" in line:
                # Tracked change line
                p = doc.add_paragraph()
                # Parse ~~deleted~~ **→** __inserted__
                parts = re.split(r"(~~.*?~~|\*\*→\*\*|__.*?__)", line)
                for part in parts:
                    if part.startswith("~~") and part.endswith("~~"):
                        run = p.add_run(part[2:-2])
                        run.font.strike = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    elif part == "**→**":
                        run = p.add_run(" → ")
                        run.bold = True
                    elif part.startswith("__") and part.endswith("__"):
                        run = p.add_run(part[2:-2])
                        run.underline = True
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif part:
                        p.add_run(part)
                continue

            # Headers
            if line_stripped.startswith("# "):
                doc.add_heading(line_stripped[2:], level=1)
            elif line_stripped.startswith("## "):
                doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith("### "):
                doc.add_heading(line_stripped[4:], level=3)
            else:
                p = doc.add_paragraph(line_stripped)

        doc.save(str(output_path))


# ─── CLI entry point ─────────────────────────────────────────────────────────


def _cli() -> None:
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate tracked-changes manuscript for peer review.",
    )
    parser.add_argument("manuscript", help="Original manuscript markdown file")
    parser.add_argument(
        "--changes", "-c",
        required=True,
        help="JSON file with changes list",
    )
    parser.add_argument(
        "-o", "--output-dir",
        default="revised",
        help="Output directory (default: revised/)",
    )
    parser.add_argument(
        "--format",
        choices=["both", "tracked", "clean"],
        default="both",
        help="Output format (default: both)",
    )
    parser.add_argument(
        "--tolerance",
        type=float,
        default=0.8,
        help="Fuzzy matching tolerance 0.0-1.0 (default: 0.8)",
    )
    parser.add_argument(
        "--manifest-only",
        action="store_true",
        help="Only generate change manifest",
    )

    args = parser.parse_args()

    # Load changes
    changes = json.loads(Path(args.changes).read_text(encoding="utf-8"))
    if isinstance(changes, dict):
        changes = changes.get("changes", changes.get("comments", []))

    gen = TrackedChangeGenerator(tolerance=args.tolerance)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Processing {len(changes)} changes...")

    if args.manifest_only:
        manifest = gen.generate_manifest(changes, output_dir / "change_manifest.md")
        print(manifest)
        return

    manuscript_name = Path(args.manuscript).stem

    if args.format in ("both", "tracked"):
        tracked_path = output_dir / f"{manuscript_name}_tracked.md"
        tracked = gen.generate_tracked(args.manuscript, changes, tracked_path)
        print(f"Tracked version: {tracked_path}")

        # Also generate .docx tracked version
        try:
            docx_tracked = output_dir / f"{manuscript_name}_tracked.docx"
            gen.export_docx(args.manuscript, changes, docx_tracked, mode="tracked")
            print(f"Tracked .docx: {docx_tracked}")
        except Exception as e:
            sys.stderr.write(f"Warning: Could not generate tracked .docx: {e}\n")

    if args.format in ("both", "clean"):
        clean_path = output_dir / f"{manuscript_name}_revised.md"
        gen.generate_clean(args.manuscript, changes, clean_path)
        print(f"Clean revised: {clean_path}")

        # Also generate .docx clean version
        try:
            docx_clean = output_dir / f"{manuscript_name}_revised.docx"
            gen.export_docx(args.manuscript, changes, docx_clean, mode="clean")
            print(f"Revised .docx: {docx_clean}")
        except Exception as e:
            sys.stderr.write(f"Warning: Could not generate clean .docx: {e}\n")

    # Always generate manifest
    manifest_path = output_dir / "change_manifest.md"
    gen.generate_manifest(changes, manifest_path)
    print(f"Change manifest: {manifest_path}")

    # Summary
    applied = sum(
        1 for c_data in changes
        if (c_data.get("status") if isinstance(c_data, dict) else c_data.status) == "applied"
    )
    print(f"\nSummary: {applied}/{len(changes)} changes applied successfully.")


if __name__ == "__main__":
    _cli()
