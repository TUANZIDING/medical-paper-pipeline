#!/usr/bin/env python3
"""
journal_formatter.py
====================
Journal-specific manuscript formatting for the medical-paper-pipeline skill.

Loads journal specifications from templates/journal_specs.json and applies:
  - Citation style (Vancouver / AMA / ICMJE / Lancet)
  - Reference formatting per journal
  - Figure/table dimension requirements
  - Font and spacing requirements
  - Abstract structure requirements
  - Keyword count limits
  - Section heading style

Generates:
  - Formatted .docx manuscript (via python-docx, optional)
  - Plain markdown with formatted references
  - Submission checklist

Usage (import as module):
    from journal_formatter import JournalFormatter
    fmt = JournalFormatter()
    fmt.load_journal("injury")
    fmt.format_manuscript("manuscript_final.md")
    fmt.export_docx("manuscript_injury.docx")

CLI usage:
    python journal_formatter.py manuscript_final.md --journal injury -o manuscript_formatted.docx
"""

from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


# ─── Citation Style Formatters ────────────────────────────────────────────────


class CitationFormatter:
    """Base class for citation style formatters."""

    def format_reference(self, ref: dict[str, Any]) -> str:
        raise NotImplementedError

    def format_in_text(self, ref: dict[str, Any], position: str = "") -> str:
        raise NotImplementedError


class VancouverFormatter(CitationFormatter):
    """
    Vancouver style: numbered citations, numbered reference list.
    Example: (1) Smith JA, et al. J Trauma. 2020;90(3):456-460.
    """

    def format_reference(self, ref: dict[str, Any]) -> str:
        """Format a reference in Vancouver numbered style."""
        parts = []

        # Authors (up to 6, then et al.)
        authors = ref.get("authors", [])
        if authors:
            if len(authors) <= 6:
                author_str = ", ".join(authors)
            else:
                author_str = ", ".join(authors[:6]) + ", et al."
            parts.append(author_str + ".")

        # Title
        if ref.get("title"):
            parts.append(f"{ref['title']}.")

        # Journal (italicized in print, here just text)
        if ref.get("journal"):
            journal = ref["journal"]
            # Remove period if title already ends with period
            if parts and not parts[-1].endswith("."):
                parts[-1] = parts[-1].rstrip(".")
            parts.append(f"{journal}.")

        # Year, volume, issue, pages
        year = ref.get("year", "")
        volume = ref.get("volume", "")
        issue = ref.get("issue", "")
        pages = ref.get("pages", "")

        if year:
            parts.append(f"{year};")
        if volume:
            parts.append(f"{volume}")
            if issue:
                parts[-1] += f"({issue})"
            if pages:
                parts[-1] += f":{pages}"
            parts[-1] += "."

        # DOI
        if ref.get("doi"):
            doi = ref["doi"]
            if not doi.startswith("10."):
                doi = "https://doi.org/" + doi
            parts.append(f"doi: {doi}")

        return " ".join(parts).strip()

    def format_in_text(self, ref: dict[str, Any], position: str = "") -> str:
        """Format an in-text citation for Vancouver style."""
        # In Vancouver, this is just a number in parentheses
        return f"({ref.get('number', '?'))"


class AMAFormatter(CitationFormatter):
    """
    AMA (American Medical Association) style: numbered, superscript or
    bracketed numbers, reference list by citation order.
    """

    def format_reference(self, ref: dict[str, Any]) -> str:
        parts = []

        authors = ref.get("authors", [])
        if authors:
            if len(authors) <= 3:
                author_str = ", ".join(authors)
            elif len(authours) == 4:
                author_str = ", ".join(authors[:4])
            else:
                author_str = f"{authors[0]}, et al."
            parts.append(author_str + ".")

        if ref.get("title"):
            parts.append(f"{ref['title']}.")

        if ref.get("journal"):
            parts.append(f"{ref['journal']}. ")

        year = ref.get("year", "")
        volume = ref.get("volume", "")
        issue = ref.get("issue", "")
        pages = ref.get("pages", "")

        if year:
            parts.append(f"{year};")
        if volume:
            parts.append(f"{volume}")
            if issue:
                parts[-1] += f"({issue})"
            if pages:
                parts[-1] += f":{pages}"
            parts[-1] += "."

        if ref.get("doi"):
            parts.append(f"doi:{ref['doi']}")

        return " ".join(parts).strip()

    def format_in_text(self, ref: dict[str, Any], position: str = "") -> str:
        return f"[{ref.get('number', '?')}]"


class ICMJEFormatter(CitationFormatter):
    """
    ICMJE (Uniform Requirements) style — same as Vancouver but with
    journal name abbreviations per NLM style.
    """

    def format_reference(self, ref: dict[str, Any]) -> str:
        # ICMJE = Vancouver + journal abbreviation
        parts = []

        authors = ref.get("authors", [])
        if authors:
            if len(authors) <= 6:
                author_str = ", ".join(authors)
            else:
                author_str = ", ".join(authors[:6]) + ", et al."
            parts.append(author_str + ".")

        if ref.get("title"):
            parts.append(ref["title"] + ".")

        # Use abbreviated journal name
        journal = ref.get("journal_abbrev", ref.get("journal", ""))
        if journal:
            parts.append(f"{journal}. ")

        year = ref.get("year", "")
        volume = ref.get("volume", "")
        issue = ref.get("issue", "")
        pages = ref.get("pages", "")

        if year:
            parts.append(f"{year};")
        if volume:
            parts.append(f"{volume}")
            if issue:
                parts[-1] += f"({issue})"
            if pages:
                parts[-1] += f":{pages}"
            parts[-1] += "."

        if ref.get("doi"):
            parts.append(f"doi:{ref['doi']}")

        return " ".join(parts).strip()

    def format_in_text(self, ref: dict[str, Any], position: str = "") -> str:
        return f"({ref.get('number', '?')})"


class LancetFormatter(CitationFormatter):
    """
    Lancet style: numbered references with superscript numbers,
    reference list with full author names, no "et al." for <6 authors.
    """

    def format_reference(self, ref: dict[str, Any]) -> str:
        parts = []

        authors = ref.get("authors", [])
        if authors:
            author_str = ", ".join(authors)
            parts.append(f"{author_str}. ")

        if ref.get("title"):
            parts.append(f"{ref['title']}. ")

        if ref.get("journal"):
            parts.append(f"{ref['journal']} ")

        year = ref.get("year", "")
        volume = ref.get("volume", "")
        issue = ref.get("issue", "")
        pages = ref.get("pages", "")

        if year:
            parts.append(f"{year};")
        if volume:
            parts.append(f"{volume}")
            if issue:
                parts[-1] += f"({issue})"
            if pages:
                parts[-1] += f":{pages}"
            parts[-1] += "."

        if ref.get("doi"):
            parts.append(f"doi:{ref['doi']}")

        return "".join(parts).strip()

    def format_in_text(self, ref: dict[str, Any], position: str = "") -> str:
        num = ref.get("number", "?")
        # Lancet uses superscript (represented as ^{n} or just the number)
        return f"^{num}"


# ─── Journal Configuration ─────────────────────────────────────────────────────


@dataclass
class JournalSpec:
    """Specification for a target journal."""

    id: str
    name: str
    full_name: str
    citation_style: str  # "vancouver" | "ama" | "icmje" | "lancet"
    abstract_limit: int = 300
    abstract_structure: list[str] = field(
        default_factory=lambda: ["background", "methods", "results", "conclusions"]
    )
    keyword_count: int = 5
    figure_width_cm: float = 8.5
    figure_width_full_cm: float = 17.0
    font: str = "Times New Roman"
    font_size: int = 12
    line_spacing: float = 1.5
    reference_limit: int = 50
    reference_format: str = "numbered"
    title_page_required: bool = True
    running_head_required: bool = False
    highlights_required: bool = False
    section_headings: dict[str, str] = field(default_factory=dict)
    # Structured abstract sections
    abstract_sections: dict[str, str] = field(default_factory=dict)
    word_count_limits: dict[str, int] = field(default_factory=dict)
    specific_requirements: list[str] = field(default_factory=list)
    submission_url: str = ""

    def get_citation_formatter(self) -> CitationFormatter:
        if self.citation_style == "ama":
            return AMAFormatter()
        elif self.citation_style == "icmje":
            return ICMJEFormatter()
        elif self.citation_style == "lancet":
            return LancetFormatter()
        else:
            return VancouverFormatter()


@dataclass
class JournalFormatter:
    """
    Formats a manuscript for a specific journal.

    Workflow:
      1. Load journal specs
      2. Apply citation style to references
      3. Apply formatting rules (font, spacing, dimensions)
      4. Generate cover letter
      5. Generate submission checklist
      6. Export to .docx (optional)
    """

    _journals: dict[str, JournalSpec] = field(default_factory=dict)
    _current_journal: JournalSpec | None = None

    def __init__(self, specs_path: str | Path | None = None):
        """
        Args:
            specs_path: Path to journal_specs.json. Defaults to
                       templates/journal_specs.json relative to the skill dir.
        """
        if specs_path is None:
            specs_path = Path(__file__).parent.parent / "templates" / "journal_specs.json"
        self._load_specs(specs_path)

    def _load_specs(self, path: str | Path) -> None:
        """Load journal specifications from JSON file."""
        p = Path(path)
        if not p.exists():
            # Use built-in defaults
            self._journals = self._default_journals()
            return

        with open(p, encoding="utf-8") as f:
            data = json.load(f)

        self._journals = {}
        for jid, spec in data.get("journals", {}).items():
            self._journals[jid] = JournalSpec(
                id=jid,
                name=spec.get("name", jid),
                full_name=spec.get("full_name", ""),
                citation_style=spec.get("citation_style", "vancouver"),
                abstract_limit=spec.get("abstract_limit", 300),
                abstract_structure=spec.get("abstract_structure", []),
                keyword_count=spec.get("keyword_count", 5),
                figure_width_cm=spec.get("figure_width_cm", 8.5),
                figure_width_full_cm=spec.get("figure_width_full_cm", 17.0),
                font=spec.get("font", "Times New Roman"),
                font_size=spec.get("font_size", 12),
                line_spacing=spec.get("line_spacing", 1.5),
                reference_limit=spec.get("reference_limit", 50),
                reference_format=spec.get("reference_format", "numbered"),
                title_page_required=spec.get("title_page_required", True),
                running_head_required=spec.get("running_head_required", False),
                highlights_required=spec.get("highlights_required", False),
                section_headings=spec.get("section_headings", {}),
                abstract_sections=spec.get("abstract_sections", {}),
                word_count_limits=spec.get("word_count_limits", {}),
                specific_requirements=spec.get("specific_requirements", []),
                submission_url=spec.get("submission_url", ""),
            )

    def list_journals(self) -> list[str]:
        """Return list of available journal IDs."""
        return list(self._journals.keys())

    def get_journal(self, journal_id: str) -> JournalSpec | None:
        """Get a journal specification by ID."""
        return self._journals.get(journal_id.lower())

    def load_journal(self, journal_id: str) -> "JournalFormatter":
        """Set the current target journal."""
        spec = self._journals.get(journal_id.lower())
        if spec is None:
            raise ValueError(
                f"Unknown journal: {journal_id}. "
                f"Available: {', '.join(self.list_journals())}"
            )
        self._current_journal = spec
        return self

    def get_current_journal(self) -> JournalSpec:
        """Get the current journal spec (raises if not set)."""
        if self._current_journal is None:
            raise ValueError("No journal selected. Call load_journal() first.")
        return self._current_journal

    # ─── Formatting ─────────────────────────────────────────────────────────

    def format_references(
        self, references: list[dict[str, Any]]
    ) -> list[str]:
        """Format a list of references for the current journal's citation style."""
        j = self.get_current_journal()
        formatter = j.get_citation_formatter()

        formatted = []
        for i, ref in enumerate(references, 1):
            ref = dict(ref)  # copy
            ref["number"] = i
            formatted.append(formatter.format_reference(ref))

        return formatted

    def format_abstract(
        self,
        abstract: str,
        sections: dict[str, str] | None = None,
    ) -> str:
        """
        Format abstract text, enforcing structure and word limit.

        Args:
            abstract: Full abstract text (may contain section headers)
            sections: Dict of section_name → section_text for structured abstracts
        """
        j = self.get_current_journal()

        if sections:
            # Structured abstract
            lines = []
            for section in j.abstract_structure:
                section_text = sections.get(section, "").strip()
                if not section_text:
                    continue
                # Check word limit per section if specified
                section_limit = j.abstract_sections.get(section, j.abstract_limit)
                words = section_text.split()
                if len(words) > section_limit:
                    section_text = " ".join(words[:section_limit]) + "..."
                lines.append(f"**{section.capitalize()}:** {section_text}")
            return "\n\n".join(lines)
        else:
            # Free-form abstract — check total limit
            words = abstract.split()
            if len(words) > j.abstract_limit:
                abstract = " ".join(words[:j.abstract_limit]) + "..."
            return abstract

    def check_keywords(self, keywords: list[str]) -> tuple[bool, list[str]]:
        """Check keyword list against journal requirements."""
        j = self.get_current_journal()
        required = j.keyword_count
        actual = len(keywords)

        if actual < required:
            return False, [f"Need {required} keywords, got {actual}"]
        if actual > required:
            return False, [f"Too many keywords: {actual} (max {required})"]
        return True, []

    def generate_submission_checklist(self) -> str:
        """Generate a journal-specific submission checklist."""
        j = self.get_current_journal()

        lines = [
            f"# Submission Checklist: {j.full_name}\n",
            f"**Citation style:** {j.citation_style.upper()}\n",
            f"**Abstract limit:** {j.abstract_limit} words\n",
            f"**Keywords:** {j.keyword_count} required\n",
            f"**Figure dimensions:** half={j.figure_width_cm}cm, full={j.figure_width_full_cm}cm\n",
            f"**Font:** {j.font} {j.font_size}pt\n",
            f"**Line spacing:** {j.line_spacing}\n",
            "\n## Pre-submission Items\n",
        ]

        # Generic checklist
        items = [
            "Manuscript file (.docx) with all figures embedded or supplied separately",
            "Abstract within word limit",
            "Keywords within limit",
            "All figures: TIFF @ 300dpi, correct dimensions",
            "All tables: editable format",
            "Reference list: reformatted per journal style",
            "In-text citations: renumbered to match reference list",
            "Cover letter drafted",
            "Title page: title + authors + affiliations + correspondence",
            "Running head (if required)",
            "Highlights/online supplementary (if required)",
            "ICMJE conflict of interest form signed",
            "Ethics statement included in Methods",
            "Funding statement included",
            "All authors have approved the manuscript",
            "No simultaneous submissions",
        ]

        for item in items:
            lines.append(f"- [ ] {item}")

        if j.specific_requirements:
            lines.extend(["\n## Journal-Specific Requirements\n"])
            for req in j.specific_requirements:
                lines.append(f"- [ ] {req}")

        if j.submission_url:
            lines.append(f"\n**Submission URL:** {j.submission_url}\n")

        return "\n".join(lines)

    def generate_cover_letter(
        self,
        title: str,
        journal_name: str,
        key_finding: str,
        novelty: str,
        corresponding_author: dict[str, str],
        ethics_statement: str = "",
        trial_registration: str = "",
        additional_notes: str = "",
    ) -> str:
        """Generate a cover letter template for the current journal."""
        j = self.get_current_journal()

        lines = [
            f"Date: [YYYY-MM-DD]\n",
            f"\nDear Editor,\n",
            f"\nWe would like to submit our manuscript titled \"{title}\" "
            f"for consideration as an original research article in {journal_name}.\n",
            f"\n**Key finding:** {key_finding}\n",
            f"\n**Novelty:** {novelty}\n",
            "\nThis manuscript reports original research that has not been published "
            "elsewhere and is not under simultaneous consideration by any other journal.\n",
        ]

        if ethics_statement:
            lines.append(f"\n**Ethical statement:** {ethics_statement}\n")

        if trial_registration:
            lines.append(f"\n**Trial registration:** {trial_registration}\n")

        if additional_notes:
            lines.append(f"\n{additional_notes}\n")

        lines.extend(
            [
                "\nWe believe this work is well-suited for {journal_name} due to its alignment "
                "with your journal's focus on [specific scope element].\n",
                "\nAll authors have approved the manuscript and declare no conflict of interest.\n",
                "\n**Corresponding author:**\n",
                f"  Name: {corresponding_author.get('name', '[Corresponding author name]')}\n",
                f"  Institution: {corresponding_author.get('institution', '[Institution]')}\n",
                f"  Email: {corresponding_author.get('email', '[email@example.com]')}\n",
                f"  Phone: {corresponding_author.get('phone', '[Phone number]')}\n",
                "\nThank you for considering our manuscript.\n",
                "\nSincerely,\n",
                f"[{corresponding_author.get('name', 'Corresponding author name')}]\n",
            ]
        )

        return "".join(lines)

    def format_manuscript(
        self,
        manuscript_md: str,
        references: list[dict[str, Any]] | None = None,
    ) -> str:
        """
        Apply journal-specific formatting to a manuscript markdown.

        Steps:
          1. Enforce abstract word limit
          2. Format reference list in journal style
          3. Re-number citations throughout the text
          4. Add journal-specific section headings
          5. Annotate with formatting notes
        """
        j = self.get_current_journal()
        text = manuscript_md

        # Format references
        if references:
            formatted_refs = self.format_references(references)
            # Replace the reference section
            ref_section = re.search(
                r"(?:^|\n)# References\n(.*?)(?:\n\n#|\Z)",
                text,
                re.DOTALL | re.IGNORECASE,
            )
            if ref_section:
                new_ref_section = "\n# References\n\n"
                for i, ref in enumerate(formatted_refs, 1):
                    new_ref_section += f"{i}. {ref}\n\n"
                text = text[: ref_section.start()] + new_ref_section + text[ref_section.end() :]

        # Enforce abstract word limit
        abstract_match = re.search(
            r"(?:^|\n)##? Abstract\n(.*?)(?:\n\n##)",
            text,
            re.DOTALL | re.IGNORECASE,
        )
        if abstract_match:
            abstract_text = abstract_match.group(1).strip()
            words = abstract_text.split()
            if len(words) > j.abstract_limit:
                # Truncate to limit
                truncated = " ".join(words[:j.abstract_limit])
                text = text[: abstract_match.start(1)] + truncated + text[abstract_match.end(1) :]
            # Add annotation
            annotation = f"\n<!-- Abstract word count: {len(words)} / limit: {j.abstract_limit} -->\n"
            text = text[: abstract_match.end(1)] + annotation + text[abstract_match.end(1) :]

        # Apply journal-specific section headings
        if j.section_headings:
            for old_heading, new_heading in j.section_headings.items():
                text = re.sub(
                    rf"(?:^|\n)(#+\s*){re.escape(old_heading)}",
                    rf"\1{new_heading}",
                    text,
                    flags=re.MULTILINE,
                )

        # Add formatting summary at top
        summary = (
            f"\n<!-- Journal: {j.full_name} | "
            f"Citation: {j.citation_style.upper()} | "
            f"Abstract limit: {j.abstract_limit} words | "
            f"Keywords: max {j.keyword_count} -->\n"
        )

        # Insert after title
        title_match = re.search(r"(?:^|\n)# .+", text)
        if title_match:
            insert_pos = title_match.end()
            text = text[:insert_pos] + summary + text[insert_pos:]

        return text

    # ─── Export ──────────────────────────────────────────────────────────────

    def export_docx(
        self,
        manuscript_md: str,
        output_path: str | Path,
        references: list[dict[str, Any]] | None = None,
    ) -> None:
        """
        Export formatted manuscript to .docx.

        Requires python-docx. Falls back to markdown if unavailable.
        """
        try:
            import docx
            from docx import Document
            from docx.shared import Pt, Cm, Inches
            from docx.enum.text import WD_LINE_SPACING

            doc = Document()

            j = self.get_current_journal()

            # Set default font
            style = doc.styles["Normal"]
            font = style.font
            font.name = j.font
            font.size = Pt(j.font_size)

            # Set paragraph spacing
            para_format = style.paragraph_format
            para_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para_format.line_spacing = j.line_spacing

            # Parse markdown and write to docx
            lines = manuscript_md.splitlines()
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                # Skip HTML comments
                if line.startswith("<!--"):
                    continue

                # Headers
                if line.startswith("# "):
                    p = doc.add_heading(line[2:], level=1)
                    self._apply_docx_style(p, j)
                elif line.startswith("## "):
                    p = doc.add_heading(line[3:], level=2)
                    self._apply_docx_style(p, j)
                elif line.startswith("### "):
                    p = doc.add_heading(line[4:], level=3)
                    self._apply_docx_style(p, j)
                elif line.startswith("**") and line.endswith("**"):
                    # Bold paragraph (abstract section)
                    p = doc.add_paragraph(line.strip("*"))
                    run = p.runs[0]
                    run.bold = True
                elif line.startswith("> "):
                    # Blockquote (abstract)
                    p = doc.add_paragraph(line[2:], style="Quote")
                elif line.startswith("-"):
                    # Bullet list
                    doc.add_paragraph(line[1:].strip(), style="List Bullet")
                elif re.match(r"^\d+\.", line):
                    # Numbered list
                    doc.add_paragraph(line, style="List Number")
                else:
                    # Normal paragraph — handle inline bold/italic
                    self._add_formatted_paragraph(doc, line, j)

            doc.save(str(output_path))
            return

        except ImportError:
            # Fallback: save as markdown with .docx extension
            # (user can rename or convert manually)
            fallback = str(output_path).replace(".docx", ".formatted.md")
            Path(fallback).write_text(manuscript_md, encoding="utf-8")
            sys.stderr.write(
                f"python-docx not available. Formatted manuscript saved as: {fallback}\n"
                "Install with: pip install python-docx\n"
            )

    def _apply_docx_style(self, paragraph, journal: JournalSpec) -> None:
        """Apply journal font/spacing to a paragraph."""
        for run in paragraph.runs:
            run.font.name = journal.font
            run.font.size = Pt(journal.font_size)

    def _add_formatted_paragraph(
        self, doc, line: str, journal: JournalSpec
    ) -> None:
        """Add a paragraph with basic inline formatting (bold/italic)."""
        import docx
        from docx import Document

        p = doc.add_paragraph()
        # Split on ** for bold, * for italic
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = journal.font
                run.font.size = Pt(journal.font_size)
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
                run.font.name = journal.font
                run.font.size = Pt(journal.font_size)
            elif part:
                run = p.add_run(part)
                run.font.name = journal.font
                run.font.size = Pt(journal.font_size)

    # ─── Built-in journal defaults ───────────────────────────────────────────

    @staticmethod
    def _default_journals() -> dict[str, JournalSpec]:
        """Return built-in journal specifications."""
        return {
            "injury": JournalSpec(
                id="injury",
                name="Injury",
                full_name="Injury - International Journal of the Care of the Injured",
                citation_style="vancouver",
                abstract_limit=250,
                abstract_structure=["background", "methods", "results", "conclusions"],
                keyword_count=4,
                figure_width_cm=8.5,
                figure_width_full_cm=17.0,
                font="Arial",
                font_size=11,
                line_spacing=1.5,
                reference_limit=40,
                title_page_required=True,
                submission_url="https://www.editorialmanager.com/INJURY",
                specific_requirements=[
                    "STROBE checklist recommended",
                    "CONSORT diagram for trials",
                    "Authorship form required",
                ],
            ),
            "jot": JournalSpec(
                id="jot",
                name="JOT",
                full_name="Journal of Orthopaedic Trauma",
                citation_style="vancouver",
                abstract_limit=250,
                keyword_count=4,
                figure_width_cm=8.5,
                figure_width_full_cm=17.0,
                font="Times New Roman",
                font_size=12,
                line_spacing=2.0,
                reference_limit=40,
                title_page_required=True,
                submission_url="https://www.editorialmanager.com/JOT",
            ),
            "lancet": JournalSpec(
                id="lancet",
                name="The Lancet",
                full_name="The Lancet",
                citation_style="lancet",
                abstract_limit=300,
                abstract_structure=["Background", "Methods", "Findings", "Interpretation"],
                keyword_count=5,
                figure_width_cm=8.5,
                figure_width_full_cm=17.5,
                font="Helvetica",
                font_size=12,
                line_spacing=1.5,
                reference_limit=50,
                title_page_required=True,
                running_head_required=True,
                highlights_required=True,
                submission_url="https://ees.elsevier.com/thelancet",
                specific_requirements=[
                    "Summary points (5 bullet points) required",
                    "Running head required",
                    "No supplemental data policy",
                ],
            ),
            "bmj": JournalSpec(
                id="bmj",
                name="BMJ",
                full_name="The BMJ",
                citation_style="vancouver",
                abstract_limit=300,
                abstract_structure=["Objective", "Design", "Setting", "Participants", "Main outcome measures", "Results", "Conclusions"],
                keyword_count=5,
                figure_width_cm=8.5,
                figure_width_full_cm=17.0,
                font="Arial",
                font_size=11,
                line_spacing=1.5,
                reference_limit=50,
                title_page_required=True,
                submission_url="https://mc.manuscriptcentral.com/bmj",
                specific_requirements=[
                    "STROBE checklist mandatory",
                    "Patient involvement statement required",
                    "Data sharing statement required",
                ],
            ),
        }


# ─── CLI entry point ─────────────────────────────────────────────────────────


def _cli() -> None:
    import argparse

    parser = argparse.ArgumentParser(
        description="Format manuscript for a target journal.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("manuscript", help="Input manuscript markdown file")
    parser.add_argument(
        "--journal", "-j",
        required=True,
        help="Target journal ID (e.g. injury, jot, bmj, lancet)",
    )
    parser.add_argument(
        "-o", "--output",
        default="manuscript_formatted.docx",
        help="Output file path (.docx or .md)",
    )
    parser.add_argument(
        "--refs",
        help="JSON file with reference data [{title, authors, year, journal, ...}]",
    )
    parser.add_argument(
        "--list-journals",
        action="store_true",
        help="List available journals and exit",
    )
    parser.add_argument(
        "--checklist-only",
        action="store_true",
        help="Only generate submission checklist",
    )
    parser.add_argument(
        "--cover-letter",
        help="Generate cover letter template (provide title)",
    )
    parser.add_argument(
        "--specs",
        help="Path to journal_specs.json (default: templates/journal_specs.json)",
    )

    args = parser.parse_args()

    formatter = JournalFormatter(specs_path=args.specs)

    if args.list_journals:
        print("Available journals:")
        for jid, spec in formatter._journals.items():
            print(f"  {jid}: {spec.full_name} ({spec.citation_style})")
        return

    if args.checklist_only:
        formatter.load_journal(args.journal)
        checklist = formatter.generate_submission_checklist()
        print(checklist)
        if args.output:
            Path(args.output).write_text(checklist, encoding="utf-8")
            print(f"\nSaved to: {args.output}")
        return

    if args.cover_letter:
        formatter.load_journal(args.journal)
        letter = formatter.generate_cover_letter(
            title=args.cover_letter,
            journal_name=formatter.get_current_journal().full_name,
            key_finding="[1-2 sentence summary of main result]",
            novelty="[What this study adds to existing literature]",
            corresponding_author={
                "name": "[Corresponding author name]",
                "institution": "[Institution]",
                "email": "[email@example.com]",
                "phone": "[Phone]",
            },
        )
        output = args.output or "cover_letter.md"
        Path(output).write_text(letter, encoding="utf-8")
        print(f"Cover letter saved to: {output}")
        return

    # Format the manuscript
    manuscript = Path(args.manuscript).read_text(encoding="utf-8")

    references = None
    if args.refs:
        references = json.loads(Path(args.refs).read_text(encoding="utf-8"))

    formatter.load_journal(args.journal)
    formatted = formatter.format_manuscript(manuscript, references)

    if args.output.endswith(".docx"):
        formatter.export_docx(formatted, args.output, references)
        print(f"Formatted manuscript saved to: {args.output}")
    else:
        out_path = Path(args.output)
        out_path.write_text(formatted, encoding="utf-8")
        print(f"Formatted manuscript saved to: {args.output}")


if __name__ == "__main__":
    _cli()
